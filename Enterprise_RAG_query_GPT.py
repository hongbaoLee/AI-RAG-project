# Enterprise_RAG_query_GPT.py from chatgpt
# This script builds a RAG system for enterprise data using ChromaDB and Qwen-max LLM via DashScope API.
# It extracts text, tables, and images from DOCX, XLSX, PDF files, and MySQL databases, then saves them to ChromaDB.
# Finally, it performs RAG queries using the LLM.
# Required libraries: langchain, chromadb, transformers, dashscope, docx, openpyxl, pandas, pdfplumber, mysql-connector-python
# Make sure to install them via pip if not already installed.
# Note: You need to have access to the Qwen-max model via DashScope API.
# Set your DashScope API key in the environment variable DASHSCOPE_API_KEY.

# MySQL数据库连接配置在MYSQL_CONFIG中，请根据实际情况修改。
# Mysql数据库检索，采取SQLDatabase + Toolkit + LLM模式。用户问题中涉及数据库相关关键词时，触发SQL生成与查询。LLM把用户问题生成SQL，采用SQLDatabaseToolkit执行查询。
# 将查询结果和文档内容一起传给LLM生成最终答案。
# LLM部分使用DeepSeek模型，通过DashScope API调用。DeepSeek模型在SQL生成方面表现更好。
# PDF文件使用pdfplumber进行文本提取，用openCV进行预处理，然后使用Tesseract进行图片识别。
# Excel 表格新增每个sheet转为HTML存入ChromaDB的功能，方便查询时保留表格结构信息。极大提升了查询效果。
# 增加了SQL语句出错重新生成机制，如果SQL执行报错，会将错误信息反馈给LLM，重新生成SQL并执行，直到成功为止。
# 修改了rag_query函数的SQL生成错误的重试机制。 如果SQL执行报错，会将错误信息反馈给LLM，重新生成SQL并执行，可以设置重试次数。

# 修改了程序中数据库查询结果返回给用户的方式。将数据库查询结果直接格式化以后返回给用户，不再送给大模型生成答案，避免大模型输出表格不完整的问题。
# 修改了数据库查询结果的格式化显示，使用tabulate库将结果以表格形式展示，提升可读性。
# 改用Langchain的SQLDatabase进行数据库查询，返回结果类型为list of dict格式，方便tabulate格式化，   2025-10-27

# 现在需要解决的问题：
# 1、用户问题的路由机制，如何准确判断用户问题是否涉及数据库查询，从而触发SQL生成与查询。
# 2、上传新的制度文件后，文档增量更新。
#
import os
import glob
import docx
import openpyxl
import pandas as pd
import pdfplumber
import pytesseract
from PIL import Image
from typing import Optional, List
from dashscope import Generation
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings                # 加载比较慢
from langchain_chroma import Chroma
from langchain.schema import Document
from chromadb.config import Settings
import mysql.connector
from sqlalchemy import create_engine
from langchain_community.utilities import SQLDatabase
from langchain_community.agent_toolkits.sql.base import SQLDatabaseToolkit
# from langchain.chains import SQLDatabaseChain
from langchain_core.language_models.llms import LLM
from openai import OpenAI
import tabulate  # 新增，需 pip install tabulate
import ast

# 配置
DASHSCOPE_API_KEY = os.getenv("DASHSCOPE_API_KEY")
QWEN_MODEL = "qwen-max"
TEXT_TO_SQL_MODEL = "glm-4.6"    # 用于SQL生成的LLM模型名称，可替换为"deepseek-v3.2-exp", "qwen3-coder-480b-a35b-instruct"等
CHROMA_DB_DIR = "./chroma_db_corp"
EMBEDDING_MODEL_NAME = "./models/st_paraphrase-multilingual-MiniLM-L12-v2"

MYSQL_CONFIG = {
    'host': 'localhost',
    'port': 3306,
    'user': 'root',
    'password': 'mysql',
    'database': 'prestige_db'
}

def get_db_schema_info():
    conn = mysql.connector.connect(**MYSQL_CONFIG)
    cursor = conn.cursor()
    # 1. 普通表结构
    cursor.execute("SHOW TABLES")
    tables = [row[0] for row in cursor.fetchall()]
    schema_info = []
    for table in tables:
        cursor.execute(f"SHOW CREATE TABLE {table}")
        create_sql = cursor.fetchone()[1]
        schema_info.append(create_sql)
        if table == "salaries":
            cursor.execute("SELECT DISTINCT salary_month FROM salaries ORDER BY salary_month DESC LIMIT 5")
            months = [str(row[0]) for row in cursor.fetchall()]
            if months:
                schema_info.append(f"{table}.salary_month 样本值: {', '.join(months)}")
    # # 2. 视图结构
    # cursor.execute("SHOW FULL TABLES WHERE Table_type = 'VIEW'")
    # views = [row[0] for row in cursor.fetchall()]
    # for view in views:
    #     cursor.execute(f"SHOW CREATE VIEW {view}")
    #     create_view = cursor.fetchone()[1]
    #     schema_info.append(f"-- 视图 {view}:\n{create_view}")
    # # 3. 存储过程结构
    # cursor.execute("SHOW PROCEDURE STATUS WHERE Db = DATABASE()")
    # procedures = [row[1] for row in cursor.fetchall()]
    # for proc in procedures:
    #     try:
    #         cursor.execute(f"SHOW CREATE PROCEDURE {proc}")
    #         proc_create = cursor.fetchone()[2]
    #         schema_info.append(f"-- 存储过程 {proc}:\n{proc_create}")
    #     except Exception as e:
    #         schema_info.append(f"-- 存储过程 {proc}: 无法获取定义，错误: {e}")
    conn.close()
    return "\n\n".join(schema_info)


DB_SCHEMA_INFO = get_db_schema_info()


#定义用于SQL生成的LLM
class Text_to_SQL_LLM(LLM):
    model: str = TEXT_TO_SQL_MODEL # 用于SQL生成的LLM
    api_key: str = DASHSCOPE_API_KEY
    base_url: str = "https://dashscope.aliyuncs.com/compatible-mode/v1"

    @property
    def _llm_type(self) -> str:
        return TEXT_TO_SQL_MODEL

    def _call(self, prompt: str, stop: Optional[List[str]] = None, **kwargs) -> str:
        client = OpenAI(api_key=self.api_key, base_url=self.base_url)
        messages = [{"role": "user", "content": prompt}]
        completion = client.chat.completions.create(
            model=self.model,
            messages=messages,
            stream=False
        )
        result = completion.choices[0].message.content.strip()
        import re
        sql_match = re.search(r"select .*?;", result, re.IGNORECASE | re.DOTALL)
        if sql_match:
            return sql_match.group(0)
        return result

# SQLDatabase + Toolkit 初始化
db_uri = f"mysql+mysqlconnector://{MYSQL_CONFIG['user']}:{MYSQL_CONFIG['password']}@{MYSQL_CONFIG['host']}:{MYSQL_CONFIG['port']}/{MYSQL_CONFIG['database']}"
db = SQLDatabase.from_uri(db_uri)
llm = Text_to_SQL_LLM()
# 改用SQLDatabaseChain方式执行SQL查询，返回类型为list of dict，方便后续格式化显示
# toolkit = SQLDatabaseToolkit(db=db, llm=llm)
# tools = toolkit.get_tools()
# query_tool = None
# for tool in tools:
#     if tool.name == "sql_db_query":
#         query_tool = tool
#         break

embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)       # 加载耗时
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_tables_from_docx(file_path):
    doc = docx.Document(file_path)
    tables = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        tables.append(pd.DataFrame(data))
    return tables

def extract_text_from_xlsx(file_path):
    wb = openpyxl.load_workbook(file_path, data_only=True)
    texts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
            if row_text.strip():
                texts.append(row_text)
    return "\n".join(texts)

def extract_tables_from_xlsx(file_path):
    wb = openpyxl.load_workbook(file_path, data_only=True)
    tables = []
    for sheet in wb.worksheets:
        data = []
        for row in sheet.iter_rows(values_only=True):
            data.append([str(cell) if cell is not None else "" for cell in row])
        tables.append(pd.DataFrame(data))
    return tables

def extract_text_tables_images_from_pdf(file_path):
    texts = []
    tables = []
    images = []
    ocr_texts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                texts.append(page_text)
            for table in page.extract_tables():
                df = pd.DataFrame(table)
                tables.append(df)
            for img in page.images:
                cropped = page.within_bbox((img["x0"], img["top"], img["x1"], img["bottom"]))
                pil_img = cropped.to_image(resolution=300).original
                images.append(pil_img)
                ocr_result = pytesseract.image_to_string(pil_img, lang="chi_sim+eng")
                if ocr_result.strip():
                    ocr_texts.append(ocr_result.strip())
    all_text = "\n".join(texts + ocr_texts)
    return all_text, tables, images

def save_to_chroma(docs: List[Document]):
    db = Chroma.from_documents(
        docs,
        embeddings,
        persist_directory=CHROMA_DB_DIR
    )

def build_corpus_and_save(file_dir):
    import chromadb
    try:
        client = chromadb.Client(Settings(persist_directory=CHROMA_DB_DIR))
        for col_info in client.list_collections():
            col = client.get_collection(name=col_info["name"])
            items = col.get(include=["ids"])
            ids = items.get("ids", [])
            if ids:
                col.delete(ids=ids)
    except Exception as e:
        print(f"Warning: failed to clear ChromaDB before build (continuing): {e}")

    docs = []
    # DOCX处理
    for file in glob.glob(os.path.join(file_dir, "*.docx")):
        text = extract_text_from_docx(file)
        for chunk in splitter.split_text(text):
            docs.append(Document(page_content=chunk, metadata={"source": file}))
        for table in extract_tables_from_docx(file):
            docs.append(Document(page_content=table.to_csv(index=False), metadata={"source": file, "type": "table"}))
    # XLSX处理（新增：每个sheet转为HTML）
    for file in glob.glob(os.path.join(file_dir, "*.xlsx")):
        sheets = pd.read_excel(file, sheet_name=None, header=None)
        for sheet_name, df in sheets.items():
            if not df.empty:
                html_str = df.to_html(index=False)
                docs.append(Document(page_content=html_str, metadata={"source": file, "sheet": sheet_name, "type": "table_html"}))
        # 兼容原有文本和csv分块（可保留/可删）
        text = extract_text_from_xlsx(file)
        for chunk in splitter.split_text(text):
            docs.append(Document(page_content=chunk, metadata={"source": file}))
        for table in extract_tables_from_xlsx(file):
            docs.append(Document(page_content=table.to_csv(index=False), metadata={"source": file, "type": "table"}))
    # PDF处理
    for file in glob.glob(os.path.join(file_dir, "*.pdf")):
        text, tables, images = extract_text_tables_images_from_pdf(file)
        for chunk in splitter.split_text(text):
            docs.append(Document(page_content=chunk, metadata={"source": file}))
        for table in tables:
            if table is not None:
                docs.append(Document(page_content=str(table), metadata={"source": file, "type": "table"}))
    save_to_chroma(docs)

def format_sql_result(result):
    """只格式化多行，单行直接自然语言输出"""
    if isinstance(result, list) and result and isinstance(result[0], dict):
        if len(result) == 1:
            # 只返回一行时，直接自然语言描述
            row = result[0]
            # 拼接成“产品xxx，销售额yyy”这样的描述
            return "，".join([f"{k}: {v}" for k, v in row.items()])
        else:
            headers = result[0].keys()
            table = [list(row.values()) for row in result]
            return tabulate.tabulate(table, headers, tablefmt="grid", stralign="center", numalign="center")
    return str(result)

def rag_query(question: str, top_k=5):
    db_chroma = Chroma(
        persist_directory=CHROMA_DB_DIR,
        embedding_function=embeddings
    )
    need_db = any(kw in question for kw in ["工资", "员工", "部门", "销售额", "奖金", "名单", "人数", "统计"])
    db_answer = ""
    sql_error = ""
    max_retry = 1  # SQL生成出错，最多重试1次
    retry_count = 0
    # if need_db and query_tool:
    if need_db:
        sql_prompt = (
            f"数据库表结构如下：\n{DB_SCHEMA_INFO}\n\n"
            "请根据表结构和样本值生成可以直接执行的MySQL SQL语句。"
            "你必须只输出SQL语句，不要输出任何解释、代码块标记或其它内容。"
            "不要输出```sql或```等代码块标记。不要加LIMIT限制，必须查询所有数据。"
            "如果涉及日期或月份，请优先使用样本值中的日期。"
            "如有视图（VIEW）或存储过程（PROCEDURE）可用，请优先使用视图或存储过程。"
            "**你必须严格使用表结构中实际存在的字段名，不要凭空创造字段。**\n"
            f"问题：{question}"
        )
        sql_str = llm.invoke(sql_prompt)
        print("生成的SQL:", sql_str)
        # db_answer = query_tool.invoke(sql_str)
        db_answer = db._execute(sql_str)
        while (
            retry_count < max_retry and
            isinstance(db_answer, str) and
            ("error" in db_answer.lower() or "unknown column" in db_answer.lower())
        ):
            retry_count += 1
            sql_error = db_answer
            retry_prompt = (
                f"数据库表结构如下：\n{DB_SCHEMA_INFO}\n\n"
                f"上一次SQL执行报错，错误信息如下：\n{sql_error}\n"
                "请根据错误信息和表结构，重新生成正确的MySQL SQL语句。"
                "你必须只输出SQL语句，不要输出任何解释、代码块标记或其它内容。"
                "不要输出```sql或```等代码块标记。不要加LIMIT限制，必须查询所有数据。"
                "如果涉及日期或月份，请优先使用实际存在的日期字段。"
                "如有视图（VIEW）或存储过程（PROCEDURE）可用，请优先使用视图或存储过程。"
                "**你必须严格使用表结构中实际存在的字段名，不要凭空创造字段。**\n"
                f"问题：{question}"
            )
            sql_str = llm.invoke(retry_prompt)
            print("重试生成的SQL:", sql_str)
            # db_answer = query_tool.invoke(sql_str)
            db_answer = db._execute(sql_str)

        # 只返回格式化SQL结果，不再拼接无关文档，也不送大模型
        formatted = format_sql_result(db_answer)
        return formatted, []
    else:
        docs = db_chroma.similarity_search(question, k=top_k)
        context = "\n".join([doc.page_content for doc in docs])
        prompt = (
            f"根据以下内容回答问题：\n\n{context}\n\n"
            f"问题：{question}\n"
            "答案："
        )
        messages = [
            {"role": "system", "content": "你是一个企业知识问答助手，回答要简洁准确。"},
            {"role": "user", "content": prompt}
        ]
        response = Generation.call(
            api_key=DASHSCOPE_API_KEY,
            model=QWEN_MODEL,
            messages=messages,
            result_format="message",
            top_p=0.8,
            temperature=0.3,
            max_tokens=512
        )
        answer = response.output.choices[0].message.content
        refs = []
        for doc in docs:
            meta = doc.metadata
            ref = f"{os.path.basename(meta.get('source', ''))}"
            if 'page' in meta:
                ref += f" (page {meta['page']})"
            refs.append(ref)
        refs = list(set(refs))
        return answer, refs

if __name__ == "__main__":
    build_corpus_and_save("./")  # 只在文档有变动时手动运行
    while True:
        question = input("请输入您的问题（输入 quit or exit 退出）：")
        if question.strip().lower() == "quit" or question.strip().lower() == "exit":
            print("程序已退出。")
            break
        answer, refs = rag_query(question)
        print("Answer:", answer)
        if refs:
            print("参考文件及页码:", "; ".join(refs))