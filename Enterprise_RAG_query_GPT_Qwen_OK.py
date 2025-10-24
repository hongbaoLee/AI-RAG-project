# Enterprise_RAG_query_GPT.py from chatgpt
# This script builds a RAG system for enterprise data using ChromaDB and Qwen-max LLM via DashScope API.
# It extracts text, tables, and images from DOCX, XLSX, PDF files, and MySQL databases, then saves them to ChromaDB.
# Finally, it performs RAG queries using the LLM.
# Required libraries: langchain, chromadb, transformers, dashscope, docx, openpyxl, pandas, pdfplumber, mysql-connector-python
# Make sure to install them via pip if not already installed.
# Note: You need to have access to the Qwen-max model via DashScope API.
# Set your DashScope API key in the environment variable DASHSCOPE_API_KEY.
#
import os
import glob
import docx
import openpyxl
import pandas as pd
import pdfplumber
import pytesseract
from PIL import Image
from dashscope import Generation
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain.schema import Document
from typing import List
from chromadb.config import Settings
from langchain_core.language_models.llms import LLM
from typing import Optional, List
import re

# LangChain SQL imports
from langchain_experimental.sql import SQLDatabaseChain
from langchain_community.utilities import SQLDatabase

import mysql.connector

# Configurations
DASHSCOPE_API_KEY = os.getenv("DASHSCOPE_API_KEY")
QWEN_MODEL = "qwen3-max"
CHROMA_DB_DIR = "./chroma_db_corp"
EMBEDDING_MODEL_NAME = "./models/st_paraphrase-multilingual-MiniLM-L12-v2"

MYSQL_CONFIG = {
    'host': 'localhost',
    'port': 3306,
    'user': 'root',
    'password': 'mysql',
    'database': 'prestige_db'
}

# --- Qwen LLM wrapper for LangChain ---
class QwenLLM(LLM):
    model: str = QWEN_MODEL
    api_key: str = DASHSCOPE_API_KEY

    @property
    def _llm_type(self) -> str:
        return "qwen"

    def _call(self, prompt: str, stop: Optional[List[str]] = None, **kwargs) -> str:
        prompt = (
            "请注意：生成的SQL语句中，所有字段都要加上表名前缀（如 employees.employee_id）。"
            "你必须只输出可以直接执行的MySQL SQL语句，不要输出任何解释、代码块标记或其它内容。"
            "不要输出``[sql或](http://_vscodecontentref_/0)``等代码块标记。"
            "并且不要加LIMIT限制，必须查询所有数据。\n"
            + prompt
        )
        messages = [
            {"role": "user", "content": prompt}
        ]
        resp = Generation.call(
            api_key=self.api_key,
            model=self.model,
            messages=messages,
            result_format="message",
            max_tokens=512,
            temperature=0.0
        )
        result = resp.output.choices[0].message.content.strip()
        # 尝试只提取第一条SQL语句
        sql_match = re.search(r"select .*?;", result, re.IGNORECASE | re.DOTALL)
        if sql_match:
            return sql_match.group(0)
        return result

# LangChain SQLDatabase setup
db_uri = f"mysql+mysqlconnector://{MYSQL_CONFIG['user']}:{MYSQL_CONFIG['password']}@{MYSQL_CONFIG['host']}:{MYSQL_CONFIG['port']}/{MYSQL_CONFIG['database']}"
db = SQLDatabase.from_uri(db_uri)
llm = QwenLLM()  # Use Qwen as the LLM
db_chain = SQLDatabaseChain.from_llm(llm, db, verbose=False)

# Embedding model
embeddings = HuggingFaceEmbeddings(model_name=EMBEDDING_MODEL_NAME)
splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=50)

def extract_text_from_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs if para.text.strip()])

def extract_tables_from_docx(file_path):
    doc = docx.Document(file_path)
    tables = []
    for table in doc.tables:
        data = []
        for row in table.rows:
            data.append([cell.text for cell in row.cells])
        tables.append(pd.DataFrame(data))
    return tables

def extract_text_from_xlsx(file_path):
    wb = openpyxl.load_workbook(file_path, data_only=True)
    texts = []
    for sheet in wb.worksheets:
        for row in sheet.iter_rows(values_only=True):
            row_text = " | ".join([str(cell) if cell is not None else "" for cell in row])
            if row_text.strip():
                texts.append(row_text)
    return "\n".join(texts)

def extract_tables_from_xlsx(file_path):
    wb = openpyxl.load_workbook(file_path, data_only=True)
    tables = []
    for sheet in wb.worksheets:
        data = []
        for row in sheet.iter_rows(values_only=True):
            data.append([str(cell) if cell is not None else "" for cell in row])
        tables.append(pd.DataFrame(data))
    return tables

def extract_text_tables_images_from_pdf(file_path):
    texts = []
    tables = []
    images = []
    ocr_texts = []
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                texts.append(page_text)
            for table in page.extract_tables():
                df = pd.DataFrame(table)
                tables.append(df)
            for img in page.images:
                cropped = page.within_bbox((img["x0"], img["top"], img["x1"], img["bottom"]))
                pil_img = cropped.to_image(resolution=300).original
                images.append(pil_img)
                ocr_result = pytesseract.image_to_string(pil_img, lang="chi_sim+eng")
                if ocr_result.strip():
                    ocr_texts.append(ocr_result.strip())
    all_text = "\n".join(texts + ocr_texts)
    return all_text, tables, images

def save_to_chroma(docs: List[Document]):
    db = Chroma.from_documents(
        docs,
        embeddings,
        persist_directory=CHROMA_DB_DIR
    )

def build_corpus_and_save(file_dir):
    # Clear ChromaDB collections
    import chromadb
    try:
        client = chromadb.Client(Settings(persist_directory=CHROMA_DB_DIR))
        for col_info in client.list_collections():
            col = client.get_collection(name=col_info["name"])
            items = col.get(include=["ids"])
            ids = items.get("ids", [])
            if ids:
                col.delete(ids=ids)
    except Exception as e:
        print(f"Warning: failed to clear ChromaDB before build (continuing): {e}")

    docs = []
    for file in glob.glob(os.path.join(file_dir, "*.docx")):
        text = extract_text_from_docx(file)
        for chunk in splitter.split_text(text):
            docs.append(Document(page_content=chunk, metadata={"source": file}))
        for table in extract_tables_from_docx(file):
            docs.append(Document(page_content=table.to_csv(index=False), metadata={"source": file, "type": "table"}))
    for file in glob.glob(os.path.join(file_dir, "*.xlsx")):
        text = extract_text_from_xlsx(file)
        for chunk in splitter.split_text(text):
            docs.append(Document(page_content=chunk, metadata={"source": file}))
        for table in extract_tables_from_xlsx(file):
            docs.append(Document(page_content=table.to_csv(index=False), metadata={"source": file, "type": "table"}))
    for file in glob.glob(os.path.join(file_dir, "*.pdf")):
        text, tables, images = extract_text_tables_images_from_pdf(file)
        for chunk in splitter.split_text(text):
            docs.append(Document(page_content=chunk, metadata={"source": file}))
        for table in tables:
            if table is not None:
                docs.append(Document(page_content=str(table), metadata={"source": file, "type": "table"}))
    save_to_chroma(docs)

def detect_need_db(question: str) -> bool:
    """Use LLM to decide if the question needs database data."""
    prompt = (
        f"用户问题：{question}\n"
        "请判断该问题是否需要实时查询数据库中的数据？只回答 yes 或 no。"
    )
    messages = [
        {"role": "system", "content": "你是一个判断问题是否需要数据库实时查询的助手。"},
        {"role": "user", "content": prompt}
    ]
    resp = Generation.call(
        api_key=DASHSCOPE_API_KEY,
        model=QWEN_MODEL,
        messages=messages,
        result_format="message",
        max_tokens=10,
        temperature=0.0
    )
    return "yes" in resp.output.choices[0].message.content.lower()

def rag_query(question: str, top_k=5):
    db = Chroma(
        persist_directory=CHROMA_DB_DIR,
        embedding_function=embeddings
    )
    docs = db.similarity_search(question, k=top_k)
    context = "\n".join([doc.page_content for doc in docs])

    # --- Use LLM to detect if DB is needed ---
    if detect_need_db(question):
        # Use LangChain SQLDatabaseChain to get real-time DB answer
        try:
            sql_result = db_chain.invoke({"query": question})
            context += f"\n\n[数据库实时数据]\n{sql_result}"
        except Exception as e:
            context += f"\n\n[数据库查询出错]: {e}"

    prompt = f"根据以下内容回答问题：\n\n{context}\n\n问题：{question}\n答案："
    messages = [
        {"role": "system", "content": "你是一个企业知识问答助手，回答要简洁准确。"},
        {"role": "user", "content": prompt}
    ]
    response = Generation.call(
        api_key=DASHSCOPE_API_KEY,
        model=QWEN_MODEL,
        messages=messages,
        result_format="message",
        top_p=0.8,
        temperature=0.3,
        max_tokens=256
    )
    answer = response.output.choices[0].message.content

    refs = []
    for doc in docs:
        meta = doc.metadata
        ref = f"{os.path.basename(meta.get('source', ''))}"
        if 'page' in meta:
            ref += f" (page {meta['page']})"
        refs.append(ref)
    refs = list(set(refs))
    return answer, refs

if __name__ == "__main__":
    build_corpus_and_save("./")  # Set your directory

    while True:
        question = input("请输入您的问题（输入 quit 退出）：")
        if question.strip().lower() == "quit":
            print("程序已退出。")
            break
        answer, refs = rag_query(question)
        print("Answer:", answer)
        print("参考文件及页码:", "; ".join(refs))