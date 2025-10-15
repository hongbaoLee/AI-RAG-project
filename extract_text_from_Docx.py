#extract text from docx file, including paragraphs and tables.

from docx import Document

def extract_text_from_docx(file_path):
    doc = Document(file_path)
    text = []

    # Extract paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            text.append(para.text.strip())

    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip().replace('\n', ' ')
                row_text.append(cell_text)
            # Join cell texts with tabs for clarity
            text.append('\t'.join(row_text))

    return "\n".join(text)

if __name__ == "__main__":
  #  file_path = "expense_rules.docx"
    file_path = "查勘费用报销管理办法.docx"
    extracted_text = extract_text_from_docx(file_path)
    print(extracted_text)

